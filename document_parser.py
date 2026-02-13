"""
Document parser module.
Extracts text from PDF, DOCX, TXT with logistics-specific optimizations:
1. PyMuPDF table detection for proper tabular data extraction
2. Tables converted to Markdown format for better RAG retrieval
3. Key-value line merging for split fields
4. Semantic enrichment: labels Pickup=Shipper, Drop=Consignee
"""

import io
import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument


# ---------------------------------------------------------------------------
# TABLE EXTRACTION using PyMuPDF's built-in table detection
# ---------------------------------------------------------------------------

def _extract_tables_as_markdown(page) -> list[str]:
    """
    Extract tables from a PDF page using PyMuPDF's find_tables().
    Returns list of markdown-formatted tables.
    """
    tables_md = []

    try:
        tables = page.find_tables()
        for table in tables:
            # Extract table data as list of lists
            data = table.extract()
            if not data or len(data) < 1:
                continue

            # Convert to markdown table
            md_lines = []

            # Process header row
            header = data[0]
            header_cells = [str(cell).strip() if cell else "" for cell in header]

            # Skip empty tables
            if all(not c for c in header_cells):
                continue

            md_lines.append("| " + " | ".join(header_cells) + " |")
            md_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

            # Process data rows
            for row in data[1:]:
                row_cells = [str(cell).strip() if cell else "" for cell in row]
                # Skip completely empty rows
                if any(c for c in row_cells):
                    md_lines.append("| " + " | ".join(row_cells) + " |")

            if len(md_lines) > 2:  # Has header + separator + at least one row
                tables_md.append("\n".join(md_lines))
    except Exception:
        # If table detection fails, return empty (fallback to text extraction)
        pass

    return tables_md


def _is_label_like(text: str) -> bool:
    """
    Heuristic to detect if text looks like a field label vs a data value.
    Labels are typically: short, no special chars, title-case or contains common label words.
    """
    if not text:
        return False
    text = text.strip().rstrip(":")

    # Too long to be a label
    if len(text) > 50:
        return False

    # Contains data-like patterns (numbers, currency, dates, emails)
    if re.search(r'^\$[\d,]+|^\d{2,}[/-]\d|@|^\d+\s*(lbs|kg|USD|units)', text):
        return False

    # Mostly alphabetic with spaces (typical label pattern)
    alpha_ratio = sum(1 for c in text if c.isalpha() or c.isspace()) / max(len(text), 1)
    if alpha_ratio > 0.8 and len(text.split()) <= 6:
        return True

    return False


def _tables_to_kv_pairs(tables_md: list[str]) -> list[str]:
    """
    Convert markdown tables to explicit key-value pairs for better retrieval.
    Works with any PDF - automatically detects:
    - Vertical KV tables (2 columns: label | value)
    - Horizontal tables (header row + data rows)
    """
    kv_pairs = []

    for table_md in tables_md:
        lines = table_md.strip().split("\n")
        if len(lines) < 3:
            continue

        # Parse header row
        header_line = lines[0]
        headers = [h.strip() for h in header_line.strip("|").split("|")]

        # Parse all data rows
        data_rows = []
        for line in lines[2:]:  # Skip header and separator
            cells = [c.strip() for c in line.strip("|").split("|")]
            data_rows.append(cells)

        # Detect table type for 2-column tables
        is_vertical_kv = False
        if len(headers) == 2:
            h0_is_label = _is_label_like(headers[0])
            h1_is_label = _is_label_like(headers[1])

            # If first col is label-like and second is value-like -> vertical KV
            # If both are label-like -> horizontal table (header row + data rows)
            if h0_is_label and not h1_is_label:
                is_vertical_kv = True
            # Both headers are labels -> it's a horizontal table with column headers

        if is_vertical_kv:
            # Vertical KV: each row is a key-value pair
            # Header row is also a KV pair
            if headers[0] and headers[1]:
                kv_pairs.append(f"{headers[0].rstrip(':')}: {headers[1]}")
            for cells in data_rows:
                if len(cells) >= 2:
                    key = cells[0].strip().rstrip(":")
                    value = cells[1].strip()
                    # Skip only truly empty values, keep N/A as it's meaningful
                    if key and value and value != "-":
                        kv_pairs.append(f"{key}: {value}")
        else:
            # Horizontal table: headers are column names, rows are records
            for cells in data_rows:
                for header, value in zip(headers, cells):
                    header_clean = header.strip().rstrip(":")
                    value_clean = value.strip()
                    # Skip only truly empty values, keep N/A as it's meaningful
                    if header_clean and value_clean and value_clean != "-":
                        kv_pairs.append(f"{header_clean}: {value_clean}")

    return kv_pairs


LOGISTICS_FIELD_LABELS = {
    "load id", "ship date", "delivery date", "po number", "po number pickup",
    "freight charges", "cod", "shipper", "consignee", "carrier", "carrier mc",
    "phone", "equipment", "agreed amount", "agreed amount (usd)", "size (in feet)",
    "load type", "shipping date", "shipping time", "delivery time",
    "appointment", "commodity", "weight", "quantity", "po/container no",
    "reference id", "created on", "booking date", "dispatcher", "email",
    "customer", "contact", "carrier pay", "total", "notes", "cod value",
    "driver name", "cell no", "cell no.", "truck no", "truck no.", "trailer no",
    "trailer no.", "type", "class", "# of units", "description of the commodity",
    "mailing address", "fax", "after hour contact", "3rd party billing",
    "transportation company", "s.no", "rate breakdown", "accepted by",
    "mode", "rate",
}

# ---------------------------------------------------------------------------
# STEP 1: Extract tabular key-value pairs from layout-sorted text
# ---------------------------------------------------------------------------

def _extract_tabular_kv(layout_text: str) -> dict:
    """
    Parse layout-sorted text to find key-value pairs from tabular regions.
    Returns dict of {label: value} found.
    """
    kv = {}

    for line in layout_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        # Single KV: 'Load ID              LD53657'
        single = re.match(r'^(.+?)\s{3,}(.+)$', stripped)
        if single:
            key = single.group(1).strip()
            val = single.group(2).strip()
            key_lower = key.lower().rstrip(":")

            if key_lower in LOGISTICS_FIELD_LABELS and val:
                # Filter: skip if value is itself mostly field labels (header-header line)
                val_words = re.split(r'\s{2,}', val)
                label_count = sum(1 for w in val_words if w.strip().lower().rstrip(":") in LOGISTICS_FIELD_LABELS)
                if label_count <= 1 and len(val) < 150:
                    # Take only the first value segment (before next label)
                    clean_val = re.split(r'\s{3,}', val)[0].strip()
                    if clean_val and clean_val.lower().rstrip(":") not in LOGISTICS_FIELD_LABELS:
                        kv[key_lower] = clean_val

    return kv


def _extract_tabular_rows(layout_text: str) -> list[str]:
    """
    Detect tabular header+value row patterns and merge them.
    E.g.: Header row: 'Carrier   Carrier MC   Phone   Equipment   Agreed Amount'
          Value row:  'SWIFT...  MC1685682    (618)   Flatbed     $400.00'
    Returns list of 'Key: Value' strings.
    """
    lines = layout_text.split("\n")
    merged_rows = []

    for i in range(len(lines) - 1):
        curr = lines[i]
        nxt = lines[i + 1]

        # Detect header-like line: multiple short words separated by large gaps
        curr_stripped = curr.strip()
        nxt_stripped = nxt.strip()
        if not curr_stripped or not nxt_stripped:
            continue

        # Count columns by splitting on 2+ spaces
        curr_cols = re.split(r'\s{2,}', curr_stripped)
        nxt_cols = re.split(r'\s{2,}', nxt_stripped)

        # If both have similar column counts (±1) and ≥3 columns
        if (len(curr_cols) >= 3 and len(nxt_cols) >= 2 and
            abs(len(curr_cols) - len(nxt_cols)) <= 2):

            # Check if header row is mostly known labels
            label_matches = sum(
                1 for c in curr_cols
                if c.strip().lower().rstrip(":") in LOGISTICS_FIELD_LABELS
            )
            if label_matches >= 2:
                # Pair columns: header[i] -> value[i]
                for j in range(min(len(curr_cols), len(nxt_cols))):
                    h = curr_cols[j].strip()
                    v = nxt_cols[j].strip()
                    if h and v and h.lower() != v.lower():
                        merged_rows.append(f"{h}: {v}")

    return merged_rows


# ---------------------------------------------------------------------------
# STEP 2: Flow-order text extraction with KV merging
# ---------------------------------------------------------------------------

def _merge_kv_lines(text: str) -> str:
    """Merge key-value pairs split across lines."""
    lines = text.split("\n")
    merged = []
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            merged.append("")
            i += 1
            continue

        line_lower = line.lower().rstrip(":")
        is_label = line_lower in LOGISTICS_FIELD_LABELS

        if not is_label and len(line.split()) <= 3 and len(line) < 30:
            if re.match(r'^[A-Za-z][A-Za-z\s\/\#\.]+$', line) and line not in ("-", "N/A", "FTL", "LTL"):
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    next_lower = next_line.lower().rstrip(":")
                    if (next_line and
                        next_lower not in LOGISTICS_FIELD_LABELS and
                        next_line != line and
                        (re.search(r'[\d\$\@\+]', next_line) or
                         next_line in ("Collect", "Prepaid", "FTL", "LTL", "N/A", "-")) and
                        len(next_line.split()) <= 8):
                        is_label = True

        if is_label and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            next_lower = next_line.lower().rstrip(":")
            if next_line and next_lower not in LOGISTICS_FIELD_LABELS:
                separator = ": " if not line.endswith(":") else " "
                merged.append(f"{line}{separator}{next_line}")
                i += 2
                continue

        merged.append(line)
        i += 1

    return "\n".join(merged)


# ---------------------------------------------------------------------------
# STEP 3: Section detection
# ---------------------------------------------------------------------------

def _group_logistics_sections(text: str) -> str:
    section_headers = [
        "Bill of Lading", "Carrier Details", "Customer Details",
        "Rate Breakdown", "Pickup", "Drop", "Standing Instructions",
        "Special Instructions", "Shipper & Carrier Instructions",
        "Driver Details", "Carrier Rate and Load Confirmation",
        "Customer Rate and Load Confirmation", "Stops",
    ]
    lines = text.split("\n")
    result = []
    for line in lines:
        stripped = line.strip()
        if stripped in section_headers:
            result.append(f"\n---SECTION: {stripped}---")
        result.append(line)
    return "\n".join(result)


# ---------------------------------------------------------------------------
# STEP 4: Semantic enrichment (Shipper / Consignee)
# ---------------------------------------------------------------------------

def _enrich_shipper_consignee(text: str) -> str:
    lines = text.split("\n")
    enriched = []

    # --- BOL format ---
    shipper_idx = None
    consignee_idx = None
    for idx, line in enumerate(lines):
        stripped = line.strip().lower()
        if stripped == "shipper":
            shipper_idx = idx
        elif stripped.startswith("consignee") and shipper_idx is not None:
            consignee_idx = idx
            break

    if shipper_idx is not None and consignee_idx is not None:
        addr_blocks = []
        j = consignee_idx + 1
        current_block = []
        while j < len(lines):
            stripped = lines[j].strip()
            if stripped == "1." or (stripped.startswith("1.") and len(stripped) <= 3):
                if current_block:
                    addr_blocks.append(current_block)
                    current_block = []
                j += 1
                continue
            if stripped in ("3rd Party Billing", "Transportation Company", "# Of Units") or stripped.startswith("---SECTION"):
                if current_block:
                    addr_blocks.append(current_block)
                break
            if stripped:
                current_block.append(stripped)
            j += 1
        if current_block:
            addr_blocks.append(current_block)

        if len(addr_blocks) >= 2:
            shipper_addr = " ".join(addr_blocks[0])
            consignee_addr = " ".join(addr_blocks[1])
            for idx, line in enumerate(lines):
                enriched.append(line)
                if idx == consignee_idx:
                    enriched.append(f"[SHIPPER (Origin): {shipper_addr}]")
                    enriched.append(f"[CONSIGNEE (Destination): {consignee_addr}]")
            lines = enriched
            enriched = []

    # --- Rate Confirmation format ---
    has_bol_enrichment = any("[SHIPPER (Origin)" in l for l in lines)
    if not has_bol_enrichment:
        in_pickup = False
        in_drop = False
        pickup_lines = []
        drop_lines = []
        pickup_done = False
        drop_done = False

        for line in lines:
            stripped = line.strip()
            if "---SECTION: Pickup---" in stripped:
                in_pickup = True
                in_drop = False
                pickup_lines = []
                enriched.append(line)
                continue
            elif "---SECTION: Drop---" in stripped:
                if pickup_lines and not pickup_done:
                    enriched.append(f"[SHIPPER / PICKUP ORIGIN: {' '.join(pickup_lines)}]")
                    pickup_done = True
                in_pickup = False
                in_drop = True
                drop_lines = []
                enriched.append(line)
                continue
            elif stripped.startswith("---SECTION:"):
                if in_drop and drop_lines and not drop_done:
                    enriched.append(f"[CONSIGNEE / DROP DESTINATION: {' '.join(drop_lines)}]")
                    drop_done = True
                in_pickup = False
                in_drop = False
                enriched.append(line)
                continue

            if in_pickup and stripped:
                if stripped.lower().startswith("po/container") or stripped.lower().startswith("load type"):
                    in_pickup = False
                    if pickup_lines and not pickup_done:
                        enriched.append(f"[SHIPPER / PICKUP ORIGIN: {' '.join(pickup_lines)}]")
                        pickup_done = True
                elif stripped not in ("Pickup", "Drop") and not stripped.startswith("---"):
                    pickup_lines.append(stripped)

            if in_drop and stripped:
                if stripped.lower().startswith("po/container") or stripped.lower().startswith("load type"):
                    in_drop = False
                    if drop_lines and not drop_done:
                        enriched.append(f"[CONSIGNEE / DROP DESTINATION: {' '.join(drop_lines)}]")
                        drop_done = True
                elif stripped not in ("Pickup", "Drop") and not stripped.startswith("---"):
                    drop_lines.append(stripped)

            enriched.append(line)

        if in_drop and drop_lines and not drop_done:
            enriched.append(f"[CONSIGNEE / DROP DESTINATION: {' '.join(drop_lines)}]")
    else:
        enriched = list(lines)

    return "\n".join(enriched)


# ---------------------------------------------------------------------------
# STEP 5: Combine layout-extracted table data with flow-extracted text
# ---------------------------------------------------------------------------

def _build_tabular_summary(kv_pairs: dict, table_rows: list[str]) -> str:
    """Build a structured summary block from tabular data."""
    if not kv_pairs and not table_rows:
        return ""

    parts = ["[STRUCTURED DATA EXTRACTED FROM TABLES]"]

    if kv_pairs:
        for k, v in kv_pairs.items():
            parts.append(f"  {k.title()}: {v}")

    if table_rows:
        seen = set()
        for row in table_rows:
            if row not in seen:
                parts.append(f"  {row}")
                seen.add(row)

    parts.append("[END STRUCTURED DATA]")
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Main parsers
# ---------------------------------------------------------------------------

def parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using table-aware strategy:
    1. PyMuPDF table detection -> markdown tables + key-value pairs
    2. Flow-order text extraction -> section-structured text
    3. Combine both for maximum RAG retrieval quality
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    # --- Pass 1: Extract tables using PyMuPDF's table detection ---
    all_tables_md = []
    all_kv_from_tables = []

    for page in doc:
        tables_md = _extract_tables_as_markdown(page)
        all_tables_md.extend(tables_md)
        kv_pairs = _tables_to_kv_pairs(tables_md)
        all_kv_from_tables.extend(kv_pairs)

    # --- Pass 2: Layout extraction for additional tabular data (fallback) ---
    all_kv = {}
    all_table_rows = []
    for page in doc:
        layout_text = page.get_text("text", sort=True)
        kv = _extract_tabular_kv(layout_text)
        all_kv.update(kv)
        rows = _extract_tabular_rows(layout_text)
        all_table_rows.extend(rows)

    # --- Pass 3: Flow-order extraction for sections ---
    pages = []
    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"[Page {page_num + 1}]\n{text.strip()}")
    doc.close()

    flow_text = "\n\n".join(pages)

    # Post-process flow text
    flow_text = _merge_kv_lines(flow_text)
    flow_text = _group_logistics_sections(flow_text)
    flow_text = _enrich_shipper_consignee(flow_text)
    flow_text = re.sub(r'\n{3,}', '\n\n', flow_text).strip()

    # --- Build combined output ---
    parts = []

    # Add extracted tables in markdown format
    if all_tables_md:
        parts.append("[EXTRACTED TABLES]")
        for i, table_md in enumerate(all_tables_md, 1):
            parts.append(f"\nTable {i}:\n{table_md}")
        parts.append("\n[END EXTRACTED TABLES]")

    # Add key-value pairs from tables
    if all_kv_from_tables:
        parts.append("\n[TABLE KEY-VALUE PAIRS]")
        seen = set()
        for kv in all_kv_from_tables:
            if kv not in seen:
                parts.append(f"  {kv}")
                seen.add(kv)
        parts.append("[END TABLE KEY-VALUE PAIRS]")

    # Add legacy tabular summary (for fields not caught by table detection)
    tabular_summary = _build_tabular_summary(all_kv, all_table_rows)
    if tabular_summary:
        parts.append(f"\n{tabular_summary}")

    # Add flow text
    parts.append(f"\n{flow_text}")

    return "\n".join(parts).strip()


def parse_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    raw = "\n\n".join(paragraphs)
    raw = _merge_kv_lines(raw)
    raw = _group_logistics_sections(raw)
    raw = _enrich_shipper_consignee(raw)
    return re.sub(r'\n{3,}', '\n\n', raw).strip()


def parse_txt(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="replace").strip()
    text = _merge_kv_lines(text)
    text = _group_logistics_sections(text)
    text = _enrich_shipper_consignee(text)
    return re.sub(r'\n{3,}', '\n\n', text).strip()


def parse_document(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext == ".docx":
        return parse_docx(file_bytes)
    elif ext == ".txt":
        return parse_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: .pdf, .docx, .txt")